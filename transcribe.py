import speech_recognition as sr
from pydub import AudioSegment
import os
import math
import time
import sys
import threading
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from colorama import init, Fore, Style

# تلاش برای ایمپورت کتابخانه‌های اصلاح متن فارسی
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    HAS_RESHAPER = True
except ImportError:
    HAS_RESHAPER = False

# راه‌اندازی رنگ‌ها برای ترمینال
init(autoreset=True)

# تنظیمات ثابت
INPUT_FOLDER = "sot"
OUTPUT_FOLDER = "outputs"
MAX_RETRIES = 3  # تعداد تلاش مجدد در صورت خطا

def fix_rtl(text):
    """
    اصلاح نمایش متن فارسی در کنسول ویندوز.
    """
    if not HAS_RESHAPER:
        return text
    try:
        reshaped_text = arabic_reshaper.reshape(text)
        bidi_text = get_display(reshaped_text)
        return bidi_text
    except Exception:
        return text

def setup_folders():
    if not os.path.exists(INPUT_FOLDER):
        os.makedirs(INPUT_FOLDER)
        print(Fore.YELLOW + fix_rtl(f"⚠ پوشه '{INPUT_FOLDER}' ساخته شد. لطفا فایل‌های صوتی را در آن قرار دهید."))
        return False
    
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
    
    return True

def save_to_docx(text, filename):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    heading = doc.add_heading(filename, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    output_path = os.path.join(OUTPUT_FOLDER, filename + ".docx")
    doc.save(output_path)
    return output_path

def format_time(seconds):
    """تبدیل ثانیه به فرمت دقیقه:ثانیه"""
    if seconds < 0: seconds = 0
    m, s = divmod(int(seconds), 60)
    return f"{m:02d}:{s:02d}"

def progress_monitor(stop_event, total_chunks, start_time, data_provider):
    """
    این تابع در یک ترد جداگانه اجرا می‌شود و نوار وضعیت را مدام بروز می‌کند.
    data_provider: تابعی که تعداد تکه‌های تمام شده و حجم کل پردازش شده را برمی‌گرداند.
    """
    bar_length = 30
    
    while not stop_event.is_set():
        completed_chunks, total_bytes_processed = data_provider()
        
        elapsed = time.time() - start_time
        if elapsed == 0: elapsed = 0.1
        
        # محاسبه درصد
        percent = (completed_chunks / total_chunks) * 100 if total_chunks > 0 else 0
        
        # نوار گرافیکی
        filled_length = int(bar_length * completed_chunks // total_chunks)
        bar = "█" * filled_length + '-' * (bar_length - filled_length)
        
        # محاسبه سرعت (KB/s)
        speed_kb = (total_bytes_processed / 1024) / elapsed
        
        # محاسبه زمان باقی‌مانده (ETA)
        avg_time_per_chunk = elapsed / completed_chunks if completed_chunks > 0 else 0
        remaining_chunks = total_chunks - completed_chunks
        
        # اگر هنوز تکه‌ای تمام نشده، ETA را تخمین بزن یا صفر نشان بده
        if completed_chunks == 0:
             # تخمین اولیه: هر تکه حدود 15 ثانیه (صرفا جهت خالی نبودن)
             eta = remaining_chunks * 15
        else:
             eta = remaining_chunks * avg_time_per_chunk

        # چاپ وضعیت (با \r برای بازنویسی خط)
        status_str = f"\r{Fore.CYAN}Processing |{bar}| {percent:.1f}% {Fore.YELLOW}[ETA: {format_time(eta)}] {Fore.MAGENTA}[Speed: {speed_kb:.1f} KB/s]   "
        sys.stdout.write(status_str)
        sys.stdout.flush()
        
        time.sleep(0.5) # آپدیت هر نیم ثانیه

def process_audio_chunk(recognizer, audio_chunk, chunk_index, language="fa-IR"):
    chunk_filename = f"temp_chunk_{chunk_index}.wav"
    audio_chunk.export(chunk_filename, format="wav")
    
    # بدست آوردن حجم فایل برای محاسبه سرعت
    file_size = os.path.getsize(chunk_filename)
    
    text = ""
    attempt = 0
    success = False

    while attempt < MAX_RETRIES and not success:
        try:
            with sr.AudioFile(chunk_filename) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data, language=language)
                success = True
        except sr.UnknownValueError:
            success = True # متن نامفهوم، ادامه میدهیم
        except sr.RequestError:
            attempt += 1
            # چاپ خطا در خط جدید تا نوار وضعیت خراب نشود
            sys.stdout.write(f"\n{Fore.RED}" + fix_rtl(f"   ✖ تکه {chunk_index + 1}: خطای اتصال (تلاش {attempt}/{MAX_RETRIES})...") + "\n")
            time.sleep(2)
        except Exception as e:
            sys.stdout.write(f"\n{Fore.RED}" + fix_rtl(f"   ✖ خطا: {e}") + "\n")
            break 
    
    if os.path.exists(chunk_filename):
        try:
            os.remove(chunk_filename)
        except:
            pass
    
    return text, file_size

def transcribe_file(file_path, current_index, total_files):
    filename = os.path.basename(file_path)
    
    print(Fore.CYAN + "\n" + "="*60)
    info_msg = f"📂 File {current_index}/{total_files}: {filename}"
    print(Fore.CYAN + Style.BRIGHT + info_msg)
    
    try:
        print(Fore.BLUE + fix_rtl("   ⏳ در حال آماده‌سازی و برش فایل..."))
        
        if file_path.endswith('.mp3'):
            sound = AudioSegment.from_mp3(file_path)
        elif file_path.endswith('.wav'):
            sound = AudioSegment.from_wav(file_path)
        else:
            sound = AudioSegment.from_file(file_path)

        sound = sound.set_channels(1).set_frame_rate(16000)

        chunk_length_ms = 60 * 1000 
        total_length_ms = len(sound)
        chunks_count = math.ceil(total_length_ms / chunk_length_ms)
        
        full_text = []
        recognizer = sr.Recognizer()
        
        start_time_file = time.time()
        
        # متغیرهای مشترک برای ترد مانیتورینگ
        shared_data = {
            'completed_chunks': 0,
            'total_bytes': 0
        }
        
        # راه‌اندازی ترد مانیتورینگ
        stop_event = threading.Event()
        
        # لامبدا برای خواندن امن متغیرها
        def get_data():
            return shared_data['completed_chunks'], shared_data['total_bytes']

        monitor_thread = threading.Thread(target=progress_monitor, args=(stop_event, chunks_count, start_time_file, get_data))
        monitor_thread.daemon = True # با بسته شدن برنامه بسته شود
        monitor_thread.start()

        # حلقه پردازش اصلی
        for i in range(chunks_count):
            start_ms = i * chunk_length_ms
            end_ms = min((i + 1) * chunk_length_ms, total_length_ms)
            chunk = sound[start_ms:end_ms]
            
            # این تابع زمان‌بر است (بلاک می‌کند)
            chunk_text, chunk_size = process_audio_chunk(recognizer, chunk, i)
            
            if chunk_text:
                full_text.append(chunk_text)
            
            # بروزرسانی آمار برای ترد مانیتورینگ
            shared_data['completed_chunks'] += 1
            shared_data['total_bytes'] += chunk_size

        # پایان کار ترد مانیتورینگ
        stop_event.set()
        monitor_thread.join()
        sys.stdout.write("\n") # خط جدید بعد از پر شدن نوار

        final_result = " ".join(full_text)
        
        if final_result.strip():
            docx_path = save_to_docx(final_result, filename)
            print(Fore.GREEN + fix_rtl(f"✅ پردازش کامل شد."))
            print(Fore.GREEN + fix_rtl(f"📄 ذخیره شده در: {docx_path}"))
        else:
            print(Fore.RED + fix_rtl(f"⛔ متنی استخراج نشد."))

    except Exception as e:
        print(Fore.RED + fix_rtl(f"\n❌ خطا: {e}"))

def main():
    print(Fore.MAGENTA + Style.BRIGHT + """
    *************************************************
    * برنامه هوشمند تبدیل صدا به متن (Ultra)    *
    *************************************************
    """)

    if not HAS_RESHAPER:
        print(Fore.YELLOW + "نکته: برای اصلاح فونت فارسی در CMD دستور زیر را بزنید:")
        print(Fore.WHITE + "pip install arabic-reshaper python-bidi\n")

    if not setup_folders():
        return

    files = sorted([f for f in os.listdir(INPUT_FOLDER) if f.lower().endswith(('.mp3', '.wav', '.ogg', '.m4a'))])
    total_files = len(files)
    
    if not files:
        print(Fore.RED + fix_rtl(f"هیچ فایل صوتی در پوشه '{INPUT_FOLDER}' پیدا نشد."))
        return

    print(Fore.WHITE + fix_rtl(f"تعداد {total_files} فایل برای پردازش پیدا شد."))
    
    for index, file_name in enumerate(files, 1):
        file_path = os.path.join(INPUT_FOLDER, file_name)
        transcribe_file(file_path, index, total_files)

    print(Fore.MAGENTA + Style.BRIGHT + "\n" + fix_rtl("✨ تمام عملیات با موفقیت به پایان رسید."))

if __name__ == "__main__":
    main()